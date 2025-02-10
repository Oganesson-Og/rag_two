"""
DOCX Document Extractor Module
---------------------------

Specialized extractor for DOCX documents with comprehensive features.

Key Features:
- Text extraction
- Style preservation
- Table detection
- Image extraction
- Header/footer handling
- Metadata parsing
- Structure analysis

Technical Details:
- python-docx integration
- Style mapping
- Content hierarchy
- Image handling
- Table processing
- Metadata extraction
- Error management

Dependencies:
- python-docx>=0.8.11
- Pillow>=8.0.0
- typing-extensions>=4.7.0

Author: Keith Satuku
Version: 2.0.0
Created: 2025
License: MIT
"""

from docx import Document
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
from .base import BaseExtractor, ExtractorResult, DocumentContent

class DocxExtractor(BaseExtractor):
    """Handles DOCX documents including preprocessing steps."""
    
    def extract(
        self,
        content: Union[str, Path, bytes],
        options: Optional[Dict[str, bool]] = None
    ) -> ExtractorResult:
        try:
            options = options or {}
            
            # Handle different input types
            if isinstance(content, (str, Path)):
                doc = Document(content)
            else:
                from io import BytesIO
                doc = Document(BytesIO(content))

            result = {
                'content': self._extract_text(doc),
                'metadata': {
                    **self.get_metadata(),
                    **self._extract_docx_metadata(doc),
                    'content_type': 'docx'
                }
            }

            if options.get('extract_tables', True):
                result['tables'] = self._extract_tables(doc)
                
            if options.get('extract_images', True):
                result['images'] = self._extract_images(doc)
                
            if options.get('extract_headers_footers', True):
                result['headers_footers'] = self._extract_headers_footers(doc)

            return result
            
        except Exception as e:
            self.logger.error(f"DOCX extraction error: {str(e)}")
            raise

    def _extract_text(self, doc: Document) -> str:
        """Extract text content while preserving structure."""
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return self._clean_text('\n'.join(text))

    def _extract_docx_metadata(self, doc: Document) -> Dict[str, Any]:
        """Extract document metadata."""
        core_properties = doc.core_properties
        return {
            'author': core_properties.author,
            'created': core_properties.created,
            'modified': core_properties.modified,
            'title': core_properties.title,
            'revision': core_properties.revision
        }

    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract tables from document."""
        tables = []
        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                data.append(row_data)
            tables.append({
                'data': data,
                'rows': len(table.rows),
                'columns': len(table.columns)
            })
        return tables

    def _extract_images(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract embedded images."""
        images = []
        rels = doc.part.rels
        for rel in rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    images.append({
                        'data': image_data,
                        'filename': rel.target_ref.split('/')[-1]
                    })
                except Exception as e:
                    self.logger.warning(f"Failed to extract image: {str(e)}")
        return images

    def _extract_headers_footers(self, doc: Document) -> Dict[str, Any]:
        """Extract headers and footers."""
        return {
            'headers': [section.header.text for section in doc.sections if section.header],
            'footers': [section.footer.text for section in doc.sections if section.footer]
        } 